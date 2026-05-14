# -*- coding: utf-8 -*-
"""将同目录 LVDS以太网控制协议.md 导出为 Word（.docx），保留分页与表格。"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt

MD_FILE = Path(__file__).with_name("LVDS以太网控制协议.md")
OUT_FILE = Path(r"c:\Users\txzing\Desktop\LVDS板卡以太网控制协议_详细版.docx")


def set_body_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def is_table_sep(cells: list[str]) -> bool:
    if not cells:
        return False
    for c in cells:
        t = c.replace(" ", "").replace("\t", "")
        if not re.fullmatch(r":?-{3,}:?", t):
            return False
    return True


def split_table_row(line: str) -> list[str]:
    s = line.strip()
    if not s.startswith("|"):
        return []
    parts = s.split("|")
    return [p.strip() for p in parts[1:-1]]


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            if r._element.rPr is not None and r._element.rPr.rFonts is not None:
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        else:
            paragraph.add_run(part)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            t.rows[i].cells[j].text = ""
            p = t.rows[i].cells[j].paragraphs[0]
            add_inline_runs(p, cell_text)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def convert(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    set_body_font(doc)

    first_h1 = True
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")

        if "page-break-before" in line and line.strip().startswith("<div"):
            add_page_break(doc)
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = split_table_row(lines[i])
                if not is_table_sep(cells):
                    rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("# "):
            title = line[2:].strip()
            if first_h1:
                doc.add_heading(title, level=0)
                first_h1 = False
            else:
                doc.add_heading(title, level=1)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            content = re.sub(r"^\d+\.\s+", "", line)
            add_inline_runs(p, content)
            i += 1
            continue

        if line.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            content = line.lstrip()[2:]
            add_inline_runs(p, content)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line.strip())
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print("Saved:", out_path)


def main() -> None:
    md = MD_FILE
    out = OUT_FILE
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    if not md.is_file():
        sys.exit(f"Missing markdown: {md}")
    convert(md, out)


if __name__ == "__main__":
    main()
